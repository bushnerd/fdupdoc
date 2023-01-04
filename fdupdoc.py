# -*- coding: UTF-8 -*-
' fdupdoc module for finding duplicate doc'

__author__ = 'scutxd'

import datetime
import logging
import os
import re
import sys

from docx import Document

formatter = logging.Formatter('%(filename)s:%(lineno)d-%(message)s')
# formatter = logging.Formatter(
#     '%(asctime)s-%(levelname)s-%(filename)s:%(lineno)d-%(message)s')
LOG_FILE_PATH = os.path.dirname(__file__)
LOG_FILE_NAME = 'result.log'
LOG_FILE = LOG_FILE_PATH + '/' + LOG_FILE_NAME

console_handler = logging.StreamHandler()  # 输出到控制台的handler
console_handler.setFormatter(formatter)
console_handler.setLevel(logging.DEBUG)  # 设置控制台日志级别为ERROR

file_handler = logging.FileHandler(LOG_FILE, mode='w')  # 输出到文件的handler
file_handler.setFormatter(formatter)
file_handler.setLevel(logging.DEBUG)  # 设置文件日志级别为DEBUG

logging.basicConfig(level=logging.DEBUG,
                    handlers=[console_handler, file_handler])

logger = logging.getLogger("log.{module_name}".format(module_name=__name__))


def getText(wordname):
    d = Document(wordname)
    texts = []
    for para in d.paragraphs:
        texts.append(para.text)
    for table in d.tables:
        for index in range(len(table.columns)):
            for row_cell in table.row_cells(index):
                if row_cell.text != '':
                    texts.append(row_cell.text)
    return texts


def is_Chinese(word):
    for ch in word:
        if '\u4e00' <= ch <= '\u9fff':
            return True
    return False


def msplit(s, seperators=',|\.|\?|，|。|？|！'):
    return re.split(seperators, s)


def readDocx(docfile):
    logger.info('*' * 80)
    logger.info('文件{}加载中……'.format(docfile))
    t1 = datetime.datetime.now()
    paras = getText(docfile)
    segs = []
    for p in paras:
        temp = []
        for s in msplit(p):
            if len(s) > 2:
                temp.append(s.replace(' ', ""))
        if len(temp) > 0:
            segs.append(temp)
    t2 = datetime.datetime.now()
    logger.info('加载完成，用时: {}'.format(t2 - t1))
    showInfo(segs, docfile)
    return segs


def showInfo(doc, filename='filename'):
    chars = 0
    segs = 0
    for p in doc:
        for s in p:
            segs = segs + 1
            chars = chars + len(s)
    logger.info('段落数: {0:>8d} 个。'.format(len(doc)))
    logger.info('短句数: {0:>8d} 句。'.format(segs))
    logger.info('字符数: {0:>8d} 个。'.format(chars))


def compareParagraph(doc1, i, doc2, j, min_segment=5):
    """
    功能为比较两个段落的相似度，返回结果为两个段落中相同字符的长度与较短段落长度的比值。
    :param p1: 行
    :param p2: 列
    :param min_segment = 5: 最小段的长度
    """
    p1 = doc1[i]
    p2 = doc2[j]
    len1 = sum([len(s) for s in p1])
    len2 = sum([len(s) for s in p2])
    # 此处关于长度的判断可能有问题，应该改为min_segment
    # if len1 < 10 or len2 < 10:
    if len1 < min_segment or len2 < min_segment:
        return []

    list = []
    for s1 in p1:
        if len(s1) < min_segment:
            continue
        for s2 in p2:
            if len(s2) < min_segment:
                continue
            if s2 in s1:
                list.append(s2)
            elif s1 in s2:
                list.append(s1)

    # 取两个字符串的最短的一个进行比值计算
    count = sum([len(s) for s in list])
    ratio = float(count) / min(len1, len2)
    if count > 10 and ratio > 0.1:
        logger.info(' 发现相同内容 '.center(80, '*'))
        logger.info('文件1第{0:0>4d}段内容：{1}'.format(i + 1, p1))
        logger.info('文件2第{0:0>4d}段内容：{1}'.format(j + 1, p2))
        logger.info('相同内容：{}'.format(list))
        logger.info('相同字符比：{1:.2f}%\n相同字符数： {0}\n'.format(count, ratio * 100))
    return list


def check_doc(checking_doc, to_check_doc):
    t1 = datetime.datetime.now()
    for i in range(len(checking_doc)):
        if i % 100 == 0:
            logger.info('处理进行中，已处理段落 {0:>4d} (总数 {1:0>4d} ） '.format(
                i, len(checking_doc)))
        for j in range(len(to_check_doc)):
            compareParagraph(checking_doc, i, to_check_doc, j)

    t2 = datetime.datetime.now()
    logger.info('比对完成，总用时: {}'.format(t2 - t1))


if (__name__ == '__main__'):
    file_name_list = []
    for file_name in os.listdir(os.path.dirname(__file__)):
        if file_name.endswith('.docx') and not os.path.isdir(file_name):
            file_name_list.append(file_name)
    logger.debug("file_name_list:{}".format(file_name_list))

    doc_dict = {}
    for file_name in file_name_list:
        doc_dict[file_name] = readDocx(file_name)

    # checking_doc 当前正在比对的主文档
    # to_check_doc_dict 当前正在比对的副文档
    to_check_doc_dict = doc_dict.copy()

    for checking_doc_name, checking_doc_text in doc_dict.items():
        to_check_doc_dict.pop(checking_doc_name)
        for to_check_doc_name, to_check_doc_text in to_check_doc_dict.items():
            logger.info('\n\n{}<----------------->{}'.format(
                checking_doc_name, to_check_doc_name))
            check_doc(checking_doc_text, to_check_doc_text)
    pass
